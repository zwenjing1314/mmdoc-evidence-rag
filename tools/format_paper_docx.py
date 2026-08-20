from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


SOURCE = Path("/Users/zhouwenjing/Desktop/小论文写作.docx")
OUTPUT = Path("artifacts/小论文写作_排版完成.docx")


REFERENCES = [
    "[1] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//Advances in Neural Information Processing Systems. 2020.",
    "[2] ROBERTSON S, ZARAGOZA H. The probabilistic relevance framework: BM25 and beyond[J]. Foundations and Trends in Information Retrieval, 2009, 3(4): 333-389.",
    "[3] KARPUKHIN V, OGUZ B, MIN S, et al. Dense passage retrieval for open-domain question answering[C]//Proceedings of EMNLP. 2020: 6769-6781.",
    "[4] REIMERS N, GUREVYCH I. Sentence-BERT: Sentence embeddings using Siamese BERT-networks[C]//Proceedings of EMNLP-IJCNLP. 2019: 3982-3992.",
    "[5] XIAO S, LIU Z, ZHANG P, et al. C-Pack: Packaged resources to advance general Chinese embedding[EB/OL]. arXiv:2309.07597, 2023.",
    "[6] CHEN J, XIAO S, ZHANG P, et al. BGE M3-Embedding: Multi-lingual, multi-functionality, multi-granularity text embeddings through self-knowledge distillation[EB/OL]. arXiv:2402.03216, 2024.",
    "[7] GAO T, YAO X, CHEN D. SimCSE: Simple contrastive learning of sentence embeddings[C]//Proceedings of EMNLP. 2021: 6894-6910.",
    "[8] KHATTAB O, ZAHARIA M. ColBERT: Efficient and effective passage search via contextualized late interaction over BERT[C]//Proceedings of SIGIR. 2020: 39-48.",
    "[9] SANTHANAM K, KHATTAB O, SAAD-FARES J, et al. ColBERTv2: Effective and efficient retrieval via lightweight late interaction[C]//Proceedings of NAACL. 2022: 3715-3734.",
    "[10] XU Y, LI M, CUI L, et al. LayoutLM: Pre-training of text and layout for document image understanding[C]//Proceedings of KDD. 2020: 1192-1200.",
    "[11] HUANG Y, LV T, CUI L, et al. LayoutLMv3: Pre-training for document AI with unified text and image masking[C]//Proceedings of ACM Multimedia. 2022: 4083-4091.",
    "[12] APPALARAJU S, TANG Y, JIA C, et al. DocFormer: End-to-end transformer for document understanding[C]//Proceedings of ICCV. 2021: 993-1003.",
    "[13] MATHEW M, KARATZAS D, JAWAHAR C V. DocVQA: A dataset for VQA on document images[C]//Proceedings of WACV. 2021: 2200-2209.",
    "[14] KIM G, HONG T, YIM M, et al. OCR-free document understanding transformer[C]//Proceedings of ECCV. 2022: 498-517.",
    "[15] MATHEW M, KARATZAS D, JAWAHAR C V. InfographicVQA[C]//Proceedings of WACV. 2022: 1697-1706.",
    "[16] CORMACK G V, CLARKE C L A, BÜTTCHER S. Reciprocal rank fusion outperforms Condorcet and individual rank learning methods[C]//Proceedings of SIGIR. 2009: 758-759.",
    "[17] DONG K, CHANG Y, GOH X D, et al. MMDocIR: Benchmarking multi-modal retrieval for long documents[EB/OL]. arXiv:2501.08828, 2025.",
    "[18] GAO T, YEN H, YU J, et al. Enabling large language models to generate text with citations[C]//Proceedings of EMNLP. 2023.",
    "[19] ASAI A, WU Z, WANG Y, et al. Self-RAG: Learning to retrieve, generate, and critique through self-reflection[C]//Proceedings of ICLR. 2024.",
    "[20] MIN S, KRASNUKHIN D, LE BRUN C, et al. FActScore: Fine-grained atomic evaluation of factual precision in long form text generation[C]//Proceedings of EMNLP. 2023: 12076-12100.",
    "[21] MANAKUL P, LIANG A, GALLEY M. SelfCheckGPT: Zero-resource black-box hallucination detection for generative large language models[C]//Proceedings of EMNLP. 2023: 9004-9017.",
    "[22] IZACARD G, GRAVE E. Leveraging passage retrieval with generative models for open domain question answering[C]//Proceedings of EACL. 2021: 874-880.",
]


def contains_cjk(text):
    return bool(re.search(r"[\u3400-\u9fff]", text or ""))


def clear_rfonts(rpr):
    rfonts = rpr.rFonts
    if rfonts is not None:
        rpr.remove(rfonts)


def set_font(run, east_asia=None, latin="Times New Roman", size=12, bold=None, italic=None):
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    if contains_cjk(run.text):
        clear_rfonts(rpr)
        return
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, east_asia="Songti SC", latin="Times New Roman", size=9)


def has_drawing(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


def has_math(paragraph):
    return bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara"))


def append_citation(paragraph, citation):
    if citation in paragraph.text:
        return
    run = paragraph.add_run(f" {citation}")
    set_font(run, size=12)


def replace_text(paragraph, old, new):
    if old not in paragraph.text:
        return
    paragraph.text = paragraph.text.replace(old, new)


def style_paragraph(paragraph):
    text = paragraph.text.strip()
    fmt = paragraph.paragraph_format
    fmt.widow_control = True
    if not text and not has_drawing(paragraph) and not has_math(paragraph):
        return
    if has_drawing(paragraph) or has_math(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_before = Pt(3)
        fmt.space_after = Pt(3)
        fmt.first_line_indent = Pt(0)
        fmt.line_spacing = 1.0
        return
    if text == "面向中文年报问答的充分性感知多粒度证据集检索方法":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(12)
        fmt.first_line_indent = Pt(0)
        for r in paragraph.runs:
            set_font(r, east_asia="Heiti SC", latin="Times New Roman", size=16, bold=True)
        return
    if text in {"摘要", "Abstract", "参考文献"} or re.match(r"^\d+\.(?!\d)", text):
        paragraph.style = "Heading 1"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.keep_with_next = True
        for r in paragraph.runs:
            set_font(r, east_asia="Heiti SC", latin="Times New Roman", size=14, bold=True)
        return
    if re.match(r"^\d+\.\d+", text):
        paragraph.style = "Heading 2"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.keep_with_next = True
        for r in paragraph.runs:
            set_font(r, east_asia="Heiti SC", latin="Times New Roman", size=12, bold=True)
        return
    if text.startswith("图 ") or text.startswith("表 ") or text.startswith("图1") or text.startswith("表1"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_before = Pt(4)
        fmt.space_after = Pt(4)
        fmt.first_line_indent = Pt(0)
        for r in paragraph.runs:
            set_font(r, east_asia="Heiti SC", latin="Times New Roman", size=10, bold=True)
        return
    if text.startswith("关键词") or text.startswith("Keywords"):
        fmt.first_line_indent = Pt(0)
        fmt.space_after = Pt(6)
        for r in paragraph.runs:
            set_font(r, size=10.5)
        return
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = Cm(0.74)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    for r in paragraph.runs:
        set_font(r, size=12)


def main():
    doc = Document(SOURCE)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.size = Pt(12)
    clear_rfonts(normal._element.get_or_add_rPr())
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    for name, size, before, after in [("Heading 1", 14, 14, 8), ("Heading 2", 12, 10, 5), ("Heading 3", 11, 8, 4)]:
        style = styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        clear_rfonts(style._element.get_or_add_rPr())
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for p in doc.paragraphs:
        style_paragraph(p)

    # Correct stale abstract metrics before adding citations.
    for p in doc.paragraphs:
        replace_text(p, "完整方法的 Page Recall@5、Region Hit@5 和 Sufficiency Rate 均达到 0.8750", "完整方法的 Page Recall@5 和 Region Hit@5 均为 0.8750，按 Top-3 返回节点重评测的 Sufficiency Rate 为 0.6875")
        replace_text(p, "the proposed method achieves 0.8750 Page Recall@5, 0.8750 Region Hit@5, and 0.8750 Sufficiency Rate", "the proposed method achieves 0.8750 Page Recall@5, 0.8750 Region Hit@5, and 0.6875 Top-3 Sufficiency Rate")
        replace_text(p, "MMDocIR Evaluation 数据[9]", "MMDocIR Evaluation 数据[17]")
        replace_text(p, "BGE-M3[10]", "BGE-M3[6]")

    # Inline citations at the first relevant use of each research line.
    for p in doc.paragraphs:
        t = p.text
        if "检索增强生成（" in t:
            append_citation(p, "[1-4]")
        if "两类信号各有优势" in t:
            append_citation(p, "[7-9]")
        if "视觉文档理解研究表明" in t:
            append_citation(p, "[10-15]")
        if "仅有相关上下文并不能保证" in t:
            append_citation(p, "[18-22]")
        if "本文在按公司划分且冻结" in t and "MMDocIR" in t:
            append_citation(p, "[17]")
        if "BM25 等词法方法" in t:
            append_citation(p, "[2-6]")
        if "采用倒数排名融合" in t:
            append_citation(p, "[16]")
        if "MMDocIR 的 Dense 与 Hybrid" in t:
            append_citation(p, "[5-6]")
        if "MMDocIR Evaluation 数据" in t and "[17]" not in t:
            append_citation(p, "[17]")

    # Reference section.
    ref_heading = doc.add_paragraph("参考文献")
    ref_heading.style = "Heading 1"
    ref_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in ref_heading.runs:
        set_font(r, east_asia="Heiti SC", latin="Times New Roman", size=14, bold=True)
    ref_heading.paragraph_format.keep_with_next = True
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        for r in p.runs:
            set_font(r, size=10.5)

    # Tables: compact but readable, with shaded headers and centered numeric cells.
    for table in doc.tables:
        table.autofit = False
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if ri == 0:
                    set_cell_shading(cell, "E8EEF5")
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        set_font(r, size=9, bold=(ri == 0))

    # Footer with page number.
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ""
    add_page_field(fp)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
